import json
import tempfile
import zipfile
from io import BytesIO
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from PIL import Image

from core.models import AITokenLedgerEntry, AITokenSettings, AITokenWallet, Book
from .ai import (
    AISelectionCoverageError,
    AIServiceError,
    action_output_contract,
    build_word_diff,
    clean_ai_draft,
    enforce_action_length,
    extract_tiptap_headings,
    extract_tiptap_text,
    generate_draft,
    infer_custom_prompt_intent,
    intent_preview_for,
)
from .intelligence import (
    build_chapter_memory,
    build_longform_engine_state,
    build_section_index,
    build_generation_context,
    check_consistency,
    infer_book_lens,
    inspect_content,
    memory_freshness,
    next_memory_meta,
    retrieve_relevant_sections,
    selection_coverage_for,
    sentence_aware_clip,
    validate_output,
)
from .docx_export import submission_prefill_for
from .language import normalize_language_code
from .models import Manuscript


class TiptapExtractionTests(TestCase):
    def test_extracts_common_tiptap_nodes(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter One"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "A first paragraph."}]},
                {
                    "type": "bulletList",
                    "content": [
                        {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "A detail"}]}]},
                    ],
                },
            ],
        }

        self.assertIn("# Chapter One", extract_tiptap_text(content))
        self.assertIn("A first paragraph.", extract_tiptap_text(content))
        self.assertIn("- A detail", extract_tiptap_text(content))
        self.assertEqual(extract_tiptap_headings(content), [{"level": 1, "text": "Chapter One"}])

    def test_handles_empty_or_malformed_content(self):
        self.assertEqual(extract_tiptap_text({}), "")
        self.assertEqual(extract_tiptap_text(None), "")
        self.assertEqual(extract_tiptap_headings({"content": "bad"}), [])


class BookIntelligenceTests(TestCase):
    def test_clean_ai_draft_removes_thinking_and_prefaces(self):
        raw = "<think>I should continue carefully.</think>\nHere's the continuation:\nMara opened the blue door."

        self.assertEqual(clean_ai_draft(raw), "Mara opened the blue door.")

    def test_generate_draft_cleans_reasoning_leakage(self):
        class FakeClient:
            def chat(self, messages, **kwargs):
                return "Thought: I will match the style.\nThis should be hidden.\n\nMara opened the blue door."

        user = get_user_model().objects.create_user(email="clean@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Clean Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Mara waited."}]}]},
        )

        result = generate_draft(manuscript, "continue", cursor_context="Mara waited.", client=FakeClient())

        self.assertEqual(result["draft"], "Mara opened the blue door.")

    def test_action_output_contracts_are_distinct_and_bounded(self):
        selected = "Mara watched the river and waited for morning. " * 20

        rewrite = action_output_contract("rewrite", selected)
        expand = action_output_contract("expand", selected)
        summarize = action_output_contract("summarize", selected)
        outline = action_output_contract("outline", selected)

        self.assertLess(rewrite["max_tokens"], outline["max_tokens"])
        self.assertLess(rewrite["max_words"], expand["max_words"])
        self.assertLess(summarize["max_words"], rewrite["max_words"])
        self.assertIn("same scope", rewrite["instruction"])
        self.assertIn("summarize only that selection", summarize["instruction"])
        self.assertIn("Do not say 'the passage'", summarize["instruction"])

    def test_balanced_and_deep_output_contracts_scale_for_large_selection_chunks(self):
        selected = "Mara watched the river and waited for morning. " * 180

        fast = action_output_contract("rewrite", selected, "fast")
        balanced = action_output_contract("rewrite", selected, "balanced")
        deep = action_output_contract("rewrite", selected, "deep")

        self.assertLess(fast["max_words"], balanced["max_words"])
        self.assertGreater(balanced["max_words"], 700)
        self.assertGreater(deep["max_words"], 900)
        self.assertGreater(balanced["max_tokens"], fast["max_tokens"])
        self.assertGreater(deep["max_tokens"], balanced["max_tokens"])

    def test_enforce_action_length_prevents_runaway_selected_text_output(self):
        selected = "Short selected paragraph with a few details."
        runaway = "word " * 500

        clipped = enforce_action_length(runaway, "rewrite", selected)

        self.assertLessEqual(len(clipped.split()), action_output_contract("rewrite", selected)["max_words"])

    def test_local_inspection_handles_structure_voice_entities_and_malformed_content(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Paris"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "Mara met John in Paris. She said, \"We should leave in 1998.\""}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "The house was quiet, patient, and full of old letters."}]},
            ],
        }

        data = inspect_content(content)

        self.assertEqual(data["chapter_map"]["section_count"], 1)
        self.assertGreater(data["voice"]["average_sentence_words"], 0)
        self.assertIn("Mara", data["entities"]["proper_nouns"])
        self.assertIn("1998", data["entities"]["dates"])
        self.assertEqual(inspect_content(None)["chapter_map"]["total_words"], 0)

    def test_book_lens_adapts_to_genre_and_book_type(self):
        fiction = infer_book_lens({"genre": "Fantasy novel"}, "The protagonist entered the scene.")
        nonfiction = infer_book_lens({"book_type": "Business strategy guide"}, "This framework helps leaders.")
        poetry = infer_book_lens({"genre": "Poetry collection"}, "A stanza of bright rain.")

        self.assertEqual(fiction["lens"], "fiction")
        self.assertEqual(nonfiction["lens"], "business")
        self.assertEqual(poetry["lens"], "poetry")
        self.assertIn("character continuity", fiction["priorities"])
        self.assertIn("framework clarity", nonfiction["priorities"])

    @override_settings(
        REEPLS_AI_CONTEXT_POLICY="balanced",
        REEPLS_AI_MAX_INPUT_CHARS_REWRITE=9000,
        REEPLS_AI_MAX_INPUT_CHARS_CONTINUE=16000,
        REEPLS_AI_MAX_INPUT_CHARS_OUTLINE=22000,
    )
    def test_context_selector_uses_action_specific_context(self):
        manuscript = Manuscript.objects.create(
            user=get_user_model().objects.create_user(email="context@example.com", password="pass12345"),
            title="Context Book",
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Opening"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Mara walked by the river. " * 1000}]},
                ],
            },
        )

        rewrite_context = build_generation_context(manuscript, "rewrite", selected_text="Mara walked.", cursor_context="river")
        summarize_context = build_generation_context(manuscript, "summarize", selected_text="Mara walked.", cursor_context="river")
        outline_context = build_generation_context(manuscript, "outline", cursor_context="river")

        self.assertEqual(rewrite_context["usage_hint"], "small context")
        self.assertEqual(outline_context["usage_hint"], "broad context")
        self.assertIn("book_lens", rewrite_context["payload"])
        self.assertLessEqual(rewrite_context["context_summary"]["excerpt_chars"], 1400)
        self.assertEqual(summarize_context["context_summary"]["excerpt_chars"], 0)
        self.assertLess(
            rewrite_context["context_summary"]["excerpt_chars"],
            outline_context["context_summary"]["excerpt_chars"],
        )

    def test_continue_context_uses_exact_cursor_section_when_available(self):
        manuscript = Manuscript.objects.create(
            user=get_user_model().objects.create_user(email="cursor@example.com", password="pass12345"),
            title="Cursor Book",
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Introduction"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Intro river promise. " * 20}]},
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Market Scene"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Market lantern betrayal. " * 20}]},
                ],
            },
        )

        context = build_generation_context(
            manuscript,
            "continue",
            cursor_context="Market lantern betrayal appears in nearby text.",
            cursor_block_index=1,
            cursor_heading="Introduction",
        )

        self.assertEqual(context["context_summary"]["current_section"], "Introduction")
        self.assertEqual(context["payload"]["chapter_map"]["current_section"]["title"], "Introduction")

    @override_settings(
        REEPLS_AI_CONTEXT_POLICY="balanced",
        REEPLS_AI_MAX_INPUT_CHARS_REWRITE=1000,
        REEPLS_AI_DEEP_SELECTION_CHUNK_CHARS=220,
        REEPLS_AI_DEEP_MAX_SELECTION_CHARS=6000,
        REEPLS_AI_BALANCED_MAX_SELECTION_CHARS=6000,
    )
    def test_selection_coverage_blocks_fast_and_chunks_balanced_or_deep(self):
        selected = (
            "Mara watched the river until the moon rose over the trees. "
            "John waited beside the gate, holding the old letter in silence. "
        ) * 25

        clipped, stop_sentence, truncated = sentence_aware_clip(selected, 180)
        fast = selection_coverage_for("rewrite", selected, cost_mode="fast")
        balanced = selection_coverage_for("rewrite", selected, cost_mode="balanced")
        deep = selection_coverage_for("rewrite", selected, cost_mode="deep")

        self.assertTrue(truncated)
        self.assertTrue(stop_sentence.endswith("."))
        self.assertLess(len(clipped), len(selected))
        self.assertFalse(fast["allowed"])
        self.assertEqual(fast["recommended_mode"], "balanced")
        self.assertTrue(balanced["allowed"])
        self.assertTrue(balanced["chunking_available"])
        self.assertGreater(balanced["estimated_chunks"], 1)
        self.assertTrue(deep["allowed"])
        self.assertTrue(deep["chunking_available"])
        self.assertGreater(deep["estimated_chunks"], 1)

    @override_settings(REEPLS_AI_MAX_INPUT_CHARS_REWRITE=1000)
    def test_generate_draft_refuses_oversized_fast_selection_before_ai_call(self):
        class FakeClient:
            def chat(self, messages, **kwargs):
                raise AssertionError("Provider should not be called for blocked coverage.")

        user = get_user_model().objects.create_user(email="coverage@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(user=user, title="Coverage Book")
        selected = "Mara watched the river. " * 200

        with self.assertRaises(AISelectionCoverageError):
            generate_draft(manuscript, "rewrite", selected_text=selected, cost_mode="fast", client=FakeClient())

    @override_settings(
        REEPLS_AI_MAX_INPUT_CHARS_REWRITE=1000,
        REEPLS_AI_BALANCED_MAX_SELECTION_CHARS=6000,
    )
    def test_generate_draft_combines_balanced_chunks(self):
        class FakeClient:
            def __init__(self):
                self.calls = 0
                self.last_usage = {}

            def chat(self, messages, **kwargs):
                self.calls += 1
                self.last_usage = {"total_tokens": 10}
                return f"Rewritten chunk {self.calls}."

        client = FakeClient()
        user = get_user_model().objects.create_user(email="chunked@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(user=user, title="Chunked Book")
        selected = (
            "Mara watched the river until the moon rose over the trees. "
            "John waited beside the gate, holding the old letter in silence. "
        ) * 25

        result = generate_draft(manuscript, "rewrite", selected_text=selected, cost_mode="balanced", client=client)

        self.assertGreater(client.calls, 1)
        self.assertIn("Rewritten chunk 1.", result["draft"])
        self.assertEqual(result["context_summary"]["chunk_count"], client.calls)
        self.assertEqual(result["context_summary"]["provider_usage"]["total_tokens"], client.calls * 10)
        self.assertEqual(result["usage_hint"], "balanced chunked context")

    def test_safety_checker_flags_empty_and_name_heavy_output(self):
        self.assertIn("empty", validate_output("continue", "", "", {})[0])
        warnings = validate_output(
            "continue",
            "Alice met Bernard and Celia near Dorchester before Everett arrived.",
            "",
            {"proper_nouns": ["Mara"]},
        )

        self.assertTrue(any("new names" in warning for warning in warnings))

    def test_generate_draft_returns_agent_metadata_and_usage(self):
        class FakeClient:
            def chat(self, messages, **kwargs):
                return "Mara followed the river into the next quiet morning."

        user = get_user_model().objects.create_user(email="draft@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Draft Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Mara watched the river."}]}]},
            ai_entities={"proper_nouns": ["Mara"]},
        )

        result = generate_draft(manuscript, "continue", cursor_context="Mara watched the river.", client=FakeClient())

        self.assertEqual(result["usage_hint"], "chapter context")
        self.assertIn("draft", result)
        self.assertGreater(result["usage"]["request_count"], 0)
        self.assertTrue(any(step["name"] == "Safety Agent" for step in result["agent_steps"]))

    def test_retrieval_finds_relevant_prior_section(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter One"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "Mara left Paris in 1998 after meeting John by the river."}]},
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter Two"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "The village prepared for rain and a quiet harvest."}]},
            ],
        }

        sections = retrieve_relevant_sections(content, "John and Mara return to Paris")

        self.assertEqual(sections[0]["title"], "Chapter One")
        self.assertIn("Mara left Paris", sections[0]["snippet"])

    def test_section_index_chunks_long_sections_for_large_manuscripts(self):
        long_text = (
            "Mara promised John she would return to Paris before winter. "
            "The river and the blue station marked the promise. "
        ) * 180
        content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "A Long Chapter"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": long_text}]},
            ],
        }

        sections = build_section_index(content)
        retrieved = retrieve_relevant_sections(content, "John Paris winter station")

        self.assertGreater(len(sections), 1)
        self.assertLessEqual(max(len(section["snippet"]) for section in sections), 1000)
        self.assertIn("A Long Chapter", retrieved[0]["title"])

    def test_longform_engine_state_reports_readiness_and_indexing(self):
        user = get_user_model().objects.create_user(email="engine@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Engine Book",
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter One"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Mara waited by the river. " * 60}]},
                ],
            },
            ai_profile_confirmed=True,
            ai_memory_stale={"is_stale": False, "changed_words": 0, "changed_chars": 0, "reason": ""},
        )

        state = build_longform_engine_state(manuscript)

        self.assertEqual(state["target_characters"], 700000)
        self.assertGreater(state["indexed_chunks"], 0)
        self.assertIn(state["readiness"], {"ready", "ready_with_warnings"})

    def test_consistency_checker_reports_contradiction_with_source(self):
        user = get_user_model().objects.create_user(email="continuity@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Continuity Book",
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter One"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Mara was alive when John reached the river."}]},
                ],
            },
            ai_entities={"proper_nouns": ["Mara", "John"]},
        )
        retrieved = retrieve_relevant_sections(manuscript.content, "Mara and John")

        report = check_consistency("Mara was dead before John reached the river.", manuscript, action="continue", retrieved_sections=retrieved)

        self.assertEqual(report["status"], "needs_review")
        self.assertEqual(report["issues"][0]["type"], "contradiction")
        self.assertIn("Chapter One", report["issues"][0]["source"])

    def test_generate_draft_sends_action_contract_to_provider(self):
        class FakeClient:
            def __init__(self):
                self.messages = None
                self.max_tokens = None

            def chat(self, messages, **kwargs):
                self.messages = messages
                self.max_tokens = kwargs.get("max_tokens")
                return "Mara waited by the river."

        client = FakeClient()
        user = get_user_model().objects.create_user(email="contract@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Contract Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Mara watched the river."}]}]},
        )

        result = generate_draft(manuscript, "rewrite", selected_text="Mara watched the river.", client=client)
        payload = json.loads(client.messages[-1]["content"])

        self.assertEqual(result["draft"], "Mara waited by the river.")
        self.assertIn("output_contract", payload)
        self.assertIn("retrieved_context", payload)
        self.assertIn("continuity_contract", payload)
        self.assertLess(client.max_tokens, 700)

    def test_summarize_contract_is_selection_scoped_and_short(self):
        class FakeClient:
            def __init__(self):
                self.messages = None
                self.max_tokens = None

            def chat(self, messages, **kwargs):
                self.messages = messages
                self.max_tokens = kwargs.get("max_tokens")
                return "Mara waits by the river, caught between memory and the promise of morning."

        selected = "Mara watched the river and waited for morning. " * 30
        client = FakeClient()
        user = get_user_model().objects.create_user(email="summary@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Summary Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": selected}]}]},
        )

        result = generate_draft(manuscript, "summarize", selected_text=selected, client=client)
        payload = json.loads(client.messages[-1]["content"])

        self.assertEqual(result["draft"], "Mara waits by the river, caught between memory and the promise of morning.")
        self.assertEqual(payload["manuscript_excerpt"], "")
        self.assertIn("summarize only that selection", payload["output_contract"]["instruction"])
        self.assertLess(client.max_tokens, 260)

    def test_custom_prompt_infers_safe_placement(self):
        selected = "Mara waited by the river."

        replace = infer_custom_prompt_intent("Improve this and make it more tense.", selected)
        insert = infer_custom_prompt_intent("Use this idea for the next scene.", selected)
        freeform = infer_custom_prompt_intent("Write a quiet transition toward morning.", "")

        self.assertEqual(replace["placement"], "replace_selection")
        self.assertEqual(insert["placement"], "insert_at_cursor")
        self.assertEqual(freeform["placement"], "insert_at_cursor")

    def test_intent_preview_and_diff_are_deterministic(self):
        preview = intent_preview_for("custom", "Improve this.", "Mara waited.", "")
        diff = build_word_diff("Mara waited by the river.", "Mara listened by the dark river.")

        self.assertEqual(preview["label"], "Replace selection")
        self.assertTrue(any(part["type"] == "removed" for part in diff))
        self.assertTrue(any(part["type"] == "added" for part in diff))

    def test_custom_prompt_sends_user_instruction_and_returns_placement(self):
        class FakeClient:
            def __init__(self):
                self.messages = None

            def chat(self, messages, **kwargs):
                self.messages = messages
                return "Mara waited, every sound of the river sharpened by the dark."

        client = FakeClient()
        user = get_user_model().objects.create_user(email="custom@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Custom Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Mara waited by the river."}]}]},
        )

        result = generate_draft(
            manuscript,
            "custom",
            selected_text="Mara waited by the river.",
            cursor_context="Mara waited by the river.",
            user_prompt="Improve this and make it more tense.",
            client=client,
        )
        payload = json.loads(client.messages[-1]["content"])

        self.assertEqual(result["placement"], "replace_selection")
        self.assertEqual(result["intent_preview"]["label"], "Replace selection")
        self.assertTrue(result["diff_available"])
        self.assertTrue(result["suggestion_diff"])
        self.assertEqual(result["cost_mode"], "balanced")
        self.assertEqual(payload["user_prompt"], "Improve this and make it more tense.")
        self.assertEqual(payload["custom_prompt_intent"]["intent"], "custom_replace")
        self.assertIn("Follow the user's instruction", payload["output_contract"]["instruction"])

    def test_custom_prompt_requires_instruction(self):
        user = get_user_model().objects.create_user(email="emptycustom@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(user=user, title="Empty Custom")

        with self.assertRaisesMessage(Exception, "Add an instruction"):
            generate_draft(manuscript, "custom", user_prompt="")

    def test_chapter_memory_and_memory_freshness_are_local(self):
        content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Chapter One"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "Mara met John in Paris in 1998."}]},
            ],
        }

        meta = next_memory_meta({}, content, "inspect")
        chapter_memory = build_chapter_memory(content, {"open_threads": ["Mara must decide."]}, meta["version"])
        freshness = memory_freshness(meta, {"is_stale": False})

        self.assertEqual(meta["version"], 1)
        self.assertEqual(freshness["label"], "Memory v1")
        self.assertEqual(chapter_memory["sections"][0]["title"], "Chapter One")
        self.assertIn("Mara", chapter_memory["sections"][0]["characters"])

    def test_cost_mode_changes_context_budget(self):
        user = get_user_model().objects.create_user(email="cost@example.com", password="pass12345")
        manuscript = Manuscript.objects.create(
            user=user,
            title="Cost Book",
            content={"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Mara waited by the river. " * 500}]}]},
            ai_cost_mode="fast",
        )

        fast = build_generation_context(manuscript, "continue", cursor_context="Mara waited.", cost_mode="fast")
        deep = build_generation_context(manuscript, "continue", cursor_context="Mara waited.", cost_mode="deep")

        self.assertEqual(fast["context_summary"]["policy"], "fast")
        self.assertEqual(deep["context_summary"]["policy"], "deep")
        self.assertLess(fast["context_summary"]["input_chars"], deep["context_summary"]["input_chars"])


class ManuscriptAITests(TestCase):
    def setUp(self):
        User = get_user_model()
        self.user = User.objects.create_user(email="author@example.com", password="pass12345")
        self.other_user = User.objects.create_user(email="other@example.com", password="pass12345")
        self.manuscript = Manuscript.objects.create(
            user=self.user,
            title="The River House",
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Opening"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "The river remembered everything."}]},
                ],
            },
        )
        self.client.force_login(self.user)

    def make_docx_upload(self, name="imported.docx"):
        document = Document()
        document.add_heading("Imported Chapter", level=1)
        document.add_paragraph("The imported river remembered its first sentence.")
        buffer = BytesIO()
        document.save(buffer)
        return SimpleUploadedFile(
            name,
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_create_blank_manuscript_still_works(self):
        response = self.client.post(reverse("write:create"), {"title": "Blank Draft", "create_mode": "scratch"})

        manuscript = Manuscript.objects.get(title="Blank Draft")
        self.assertRedirects(response, reverse("write:editor", args=[manuscript.pk]))
        self.assertEqual(manuscript.content, {})

    def test_create_manuscript_from_docx_upload(self):
        upload = self.make_docx_upload()

        response = self.client.post(
            reverse("write:create"),
            {"title": "Imported Draft", "create_mode": "upload", "manuscript_file": upload},
        )

        manuscript = Manuscript.objects.get(title="Imported Draft")
        self.assertRedirects(response, reverse("write:editor", args=[manuscript.pk]))
        nodes = manuscript.content["content"]
        self.assertEqual(nodes[0]["type"], "heading")
        self.assertIn("Imported Chapter", nodes[0]["content"][0]["text"])
        self.assertIn("imported river", nodes[1]["content"][0]["text"])

    @patch("write.views.default_storage.save")
    def test_uploaded_source_file_is_not_saved(self, storage_save):
        response = self.client.post(
            reverse("write:create"),
            {"title": "No Storage", "create_mode": "upload", "manuscript_file": self.make_docx_upload("no-storage.docx")},
        )

        self.assertEqual(response.status_code, 302)
        storage_save.assert_not_called()

    @patch("write.imports.PdfReader")
    def test_create_manuscript_from_pdf_upload(self, reader_mock):
        class Page:
            def extract_text(self):
                return "PDF opening paragraph.\n\nSecond PDF paragraph."

        reader_mock.return_value.pages = [Page()]
        upload = SimpleUploadedFile("imported.pdf", b"%PDF-1.4", content_type="application/pdf")

        response = self.client.post(
            reverse("write:create"),
            {"title": "PDF Draft", "create_mode": "upload", "manuscript_file": upload},
        )

        manuscript = Manuscript.objects.get(title="PDF Draft")
        self.assertRedirects(response, reverse("write:editor", args=[manuscript.pk]))
        text = json.dumps(manuscript.content)
        self.assertIn("PDF opening paragraph", text)
        self.assertIn("Second PDF paragraph", text)

    def test_create_rejects_unsupported_upload_type(self):
        upload = SimpleUploadedFile("notes.txt", b"hello", content_type="text/plain")

        response = self.client.post(
            reverse("write:create"),
            {"title": "Bad Upload", "create_mode": "upload", "manuscript_file": upload},
        )

        self.assertEqual(response.status_code, 400)
        self.assertContains(response, "Only DOCX and PDF files", status_code=400)
        self.assertFalse(Manuscript.objects.filter(title="Bad Upload").exists())

    def test_create_rejects_large_upload(self):
        upload = SimpleUploadedFile(
            "large.pdf",
            b"x" * (10 * 1024 * 1024 + 1),
            content_type="application/pdf",
        )

        response = self.client.post(
            reverse("write:create"),
            {"title": "Too Large", "create_mode": "upload", "manuscript_file": upload},
        )

        self.assertEqual(response.status_code, 400)
        self.assertContains(response, "smaller than 10 MB", status_code=400)
        self.assertFalse(Manuscript.objects.filter(title="Too Large").exists())

    @patch("write.imports.PdfReader")
    def test_create_rejects_empty_extraction(self, reader_mock):
        reader_mock.return_value.pages = []
        upload = SimpleUploadedFile("empty.pdf", b"%PDF-1.4", content_type="application/pdf")

        response = self.client.post(
            reverse("write:create"),
            {"title": "Empty Import", "create_mode": "upload", "manuscript_file": upload},
        )

        self.assertEqual(response.status_code, 400)
        self.assertContains(response, "No readable text", status_code=400)
        self.assertFalse(Manuscript.objects.filter(title="Empty Import").exists())

    def test_landing_renders_delete_action_with_confirmation(self):
        response = self.client.get(reverse("write:landing"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, reverse("write:delete", args=[self.manuscript.pk]))
        self.assertContains(response, "Delete manuscript")
        self.assertContains(response, "Delete manuscript?")
        self.assertContains(response, "Delete forever")
        self.assertContains(response, "deleteTarget")
        self.assertContains(response, "This cannot be undone")
        self.assertNotContains(response, "return confirm")

    def test_delete_manuscript_removes_owned_manuscript(self):
        response = self.client.post(reverse("write:delete", args=[self.manuscript.pk]))

        self.assertRedirects(response, reverse("write:landing"))
        self.assertFalse(Manuscript.objects.filter(pk=self.manuscript.pk).exists())

    def test_delete_manuscript_is_owner_only(self):
        self.client.force_login(self.other_user)

        response = self.client.post(reverse("write:delete", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 404)
        self.assertTrue(Manuscript.objects.filter(pk=self.manuscript.pk).exists())

    def test_walkthrough_seen_endpoint_is_owner_only_and_marks_seen(self):
        response = self.client.post(reverse("write:walkthrough_seen", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 200)
        self.manuscript.refresh_from_db()
        self.assertTrue(self.manuscript.editor_walkthrough_seen)

        self.client.force_login(self.other_user)
        response = self.client.post(reverse("write:walkthrough_seen", args=[self.manuscript.pk]))
        self.assertEqual(response.status_code, 404)

    @override_settings(DEEPSEEK_API_KEY="")
    def test_missing_api_key_returns_clear_error(self):
        response = self.client.post(reverse("write:ai_analyze", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 503)
        self.assertIn("AI is not configured", response.json()["message"])

    def test_editor_renders_ai_panel(self):
        response = self.client.get(reverse("write:editor", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Reepls Editor AI")
        self.assertContains(response, "ed-floating-ai")
        self.assertContains(response, "Reepls Ai")
        self.assertContains(response, "ask, rewrite, continue")
        self.assertContains(response, 'maxlength="1000"')
        self.assertContains(response, "send")
        self.assertContains(response, "Fast")
        self.assertContains(response, "Balanced")
        self.assertContains(response, "Deep")
        self.assertContains(response, "Clean")
        self.assertContains(response, "Changes")
        self.assertContains(response, "ai-floating-suggestion")
        self.assertContains(response, "Suggestion ready")
        self.assertContains(response, "setSuggestionView")
        self.assertContains(response, "ai-mode-segment")
        self.assertContains(response, "ai-action-square")
        self.assertContains(response, "Agents working")
        self.assertContains(response, "ai-agent-loader")
        self.assertNotContains(response, "ai-suggestion-bar")
        self.assertNotContains(response, "ai-floating-suggestion-warning")
        self.assertNotContains(response, "ai-floating-suggestion-title::before")
        self.assertContains(response, "Quickly reads your chapters and writing style. Free.")
        self.assertContains(response, "Updates the deeper understanding of your book. Uses AI tokens.")
        self.assertContains(response, "https://wa.me/237682268375")
        self.assertContains(response, "WhatsApp help")
        self.assertContains(response, "capturePromptSelection")
        self.assertContains(response, "ai-prompt-selection")
        self.assertContains(response, "Status")
        self.assertContains(response, "ed-ai-profile-data")
        self.assertContains(response, "ed-ai-voice-data")
        self.assertContains(response, "ed-ai-memory-freshness-data")
        self.assertContains(response, "ed-ai-chapter-memory-data")
        self.assertContains(response, "ed-ai-longform-data")
        self.assertContains(response, "ed-ai-token-status-data")
        self.assertContains(response, "Reepls AI Tokens")
        self.assertContains(response, "Add tokens")
        self.assertContains(response, "Mobile Money (OM/MOMO)")
        self.assertContains(response, "Card")
        self.assertContains(response, reverse("write:download_docx", args=[self.manuscript.pk]))
        self.assertContains(response, reverse("write:submit_to_xanula", args=[self.manuscript.pk]))
        self.assertContains(response, "Submit to Xanula")
        self.assertNotContains(response, "ed-page-break")
        self.assertContains(response, "ed-walkthrough-data")
        self.assertContains(response, "Welcome to Reepls Editor AI")

    def test_download_manuscript_as_docx(self):
        response = self.client.get(reverse("write:download_docx", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response["Content-Type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("the-river-house.docx", response["Content-Disposition"])
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("The River House", document_xml)
        self.assertIn("The river remembered everything.", document_xml)

    def test_submission_prefill_detects_french_and_avoids_wrong_language_summary(self):
        self.manuscript.content = {
            "type": "doc",
            "content": [
                {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Ouverture"}]},
                {"type": "paragraph", "content": [{"type": "text", "text": "La riviere garde les secrets des familles. Elle raconte une histoire dans la nuit pour les enfants."}]},
            ],
        }
        self.manuscript.ai_profile = {"language": "Francais", "genre": "Roman"}
        self.manuscript.ai_memory = {"book_summary": "A story about a river and old family secrets."}

        prefill = submission_prefill_for(self.manuscript)

        self.assertEqual(prefill["language"], Book.Language.FRENCH)
        self.assertIn("La riviere", prefill["long_description"])
        self.assertNotIn("A story about", prefill["long_description"])
        self.assertEqual(normalize_language_code("anglais"), Book.Language.ENGLISH)

    def test_submit_to_xanula_prefills_publish_form_and_attaches_docx(self):
        file_storage = {
            "default": {"BACKEND": "django.core.files.storage.FileSystemStorage"},
            "staticfiles": {"BACKEND": "django.contrib.staticfiles.storage.StaticFilesStorage"},
        }
        with tempfile.TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root, STORAGES=file_storage):
            response = self.client.post(reverse("write:submit_to_xanula", args=[self.manuscript.pk]))
            self.assertEqual(response.status_code, 302)
            self.assertEqual(response.url, f"/publish/?from_manuscript={self.manuscript.pk}")

            page = self.client.get(response.url)
            self.assertContains(page, "The River House")
            self.assertContains(page, "Imported from Reepls Write")
            self.assertContains(page, "the-river-house.docx")

            image_buffer = BytesIO()
            Image.new("RGB", (16, 24), color=(120, 30, 30)).save(image_buffer, format="PNG")
            cover = SimpleUploadedFile("cover.png", image_buffer.getvalue(), content_type="image/png")
            submit_response = self.client.post(
                reverse("core:publish_book"),
                data={
                    "write_manuscript_id": str(self.manuscript.pk),
                    "title": "The River House",
                    "short_description": "A river story",
                    "long_description": "A river story with memory and mystery.",
                    "category": Book.Category.FICTION,
                    "language": Book.Language.ENGLISH,
                    "price": "0",
                    "hard_copy_option": Book.HardCopyOption.NONE,
                    "cover_image": cover,
                },
            )

            self.assertEqual(submit_response.status_code, 302)
            book = Book.objects.get(title="The River House")
            self.assertTrue(book.manuscript_file.name.endswith(".docx"))
            self.assertTrue(book.manuscript_file.storage.exists(book.manuscript_file.name))
            self.assertNotIn("write_submission_prefill", self.client.session)

    def test_only_owner_can_call_ai_endpoint(self):
        self.client.force_login(self.other_user)

        response = self.client.post(reverse("write:ai_generate", args=[self.manuscript.pk]), data={}, content_type="application/json")

        self.assertEqual(response.status_code, 404)

    @patch("write.views.analyze_manuscript")
    def test_analyze_saves_profile_and_memory(self, analyze_mock):
        analyze_mock.return_value = (
            {
                "genre": "Literary fiction",
                "tone": "Reflective",
                "target_audience": "Adults",
                "language": "English",
                "book_type": "Novel",
                "style_notes": "Lyrical but clear.",
            },
            {
                "book_summary": "A story about memory and home.",
                "chapter_summaries": [],
                "characters": ["Mara"],
                "locations": ["River House"],
                "timeline_notes": [],
                "consistency_notes": [],
            },
        )

        response = self.client.post(reverse("write:ai_analyze", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 200)
        self.manuscript.refresh_from_db()
        self.assertEqual(self.manuscript.ai_profile["genre"], "Literary fiction")
        self.assertEqual(self.manuscript.ai_memory["characters"], ["Mara"])
        self.assertEqual(self.manuscript.ai_memory_meta["version"], 1)
        self.assertTrue(self.manuscript.ai_chapter_memory["sections"])
        self.assertFalse(self.manuscript.ai_profile_confirmed)
        self.assertLess(response.json()["token_delta"], 0)
        self.assertEqual(AITokenLedgerEntry.objects.filter(user=self.user, entry_type=AITokenLedgerEntry.EntryType.USAGE).count(), 1)

    def test_profile_endpoint_saves_confirmed_edits(self):
        response = self.client.post(
            reverse("write:ai_profile", args=[self.manuscript.pk]),
            data=json.dumps({
                "profile": {
                    "genre": "Memoir",
                    "tone": "Warm",
                    "target_audience": "General readers",
                    "language": "English",
                    "book_type": "Nonfiction",
                    "style_notes": "Plainspoken.",
                },
                "confirmed": True,
            }),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.manuscript.refresh_from_db()
        self.assertEqual(self.manuscript.ai_profile["genre"], "Memoir")
        self.assertEqual(self.manuscript.ai_profile["language"], "English")
        self.assertEqual(self.manuscript.ai_memory_meta["version"], 1)
        self.assertTrue(self.manuscript.ai_profile_confirmed)

    def test_profile_endpoint_normalizes_french_language(self):
        response = self.client.post(
            reverse("write:ai_profile", args=[self.manuscript.pk]),
            data=json.dumps({
                "profile": {
                    "genre": "Roman",
                    "tone": "Lyrical",
                    "target_audience": "Adultes",
                    "language": "francais",
                    "book_type": "Roman",
                    "style_notes": "Sobre.",
                },
                "confirmed": True,
            }),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.manuscript.refresh_from_db()
        self.assertEqual(self.manuscript.ai_profile["language"], "French")

    @override_settings(DEEPSEEK_API_KEY="")
    def test_inspect_endpoint_runs_without_api_key(self):
        response = self.client.post(reverse("write:ai_inspect", args=[self.manuscript.pk]))

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["usage_hint"], "local only")
        self.assertIn("voice", data)
        self.assertIn("chapter_map", data)
        self.assertIn("entities", data)
        self.assertIn("engine_state", data)
        self.assertIn("memory_freshness", data)
        self.assertIn("chapter_memory", data)
        self.assertEqual(data["engine_state"]["retrieval"]["strategy"], "section-aware local retrieval")
        wallet = AITokenWallet.objects.get(user=self.user)
        self.assertEqual(wallet.balance, 0)
        self.assertIsNone(wallet.free_grant_at)

    def test_save_marks_memory_stale_after_large_change(self):
        response = self.client.post(
            reverse("write:save", args=[self.manuscript.pk]),
            data=json.dumps({
                "title": self.manuscript.title,
                "content": {
                    "type": "doc",
                    "content": [
                        {"type": "paragraph", "content": [{"type": "text", "text": "new words " * 120}]},
                    ],
                },
            }),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.json()["stale"]["is_stale"])
        self.manuscript.refresh_from_db()
        self.assertTrue(self.manuscript.ai_memory_stale["is_stale"])

    @patch("write.views.generate_draft")
    def test_generate_returns_draft_without_mutating_content(self, generate_mock):
        original_content = self.manuscript.content
        generate_mock.return_value = {
            "draft": "The river carried the next secret downstream.",
            "placement": "insert_at_cursor",
            "intent_preview": {"label": "Insert at cursor", "placement": "insert_at_cursor"},
            "diff_available": False,
            "suggestion_diff": [],
            "agent_steps": [{"name": "Safety Agent", "status": "ready", "detail": "Checked suggestion."}],
            "usage_hint": "chapter context",
            "safety_warnings": ["Suggestion passed local safety checks."],
            "consistency_report": {"status": "clear", "issues": [], "retrieved_sections": [], "reviewed_action": "continue"},
            "engine_state": {"readiness": "ready", "integrity_score": 100},
            "memory_freshness": {"label": "Memory v1"},
            "chapter_memory": {"sections": []},
            "cost_mode": "balanced",
            "context_summary": {"input_chars": 500},
            "usage": {"request_count": 1, "input_chars": 500, "output_chars": 45, "actions": {}, "last_request": {}},
        }

        response = self.client.post(
            reverse("write:ai_generate", args=[self.manuscript.pk]),
            data=json.dumps({"action": "continue", "selected_text": "", "cursor_context": "river"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["draft"], "The river carried the next secret downstream.")
        self.assertEqual(response.json()["placement"], "insert_at_cursor")
        self.assertEqual(response.json()["intent_preview"]["label"], "Insert at cursor")
        self.assertFalse(response.json()["diff_available"])
        self.assertEqual(response.json()["cost_mode"], "balanced")
        self.assertEqual(response.json()["usage_hint"], "chapter context")
        self.assertLess(response.json()["token_balance"], 10000)
        self.assertLess(response.json()["token_delta"], 0)
        self.assertEqual(response.json()["engine_state"]["readiness"], "ready")
        self.manuscript.refresh_from_db()
        self.assertEqual(self.manuscript.content, original_content)
        self.assertEqual(self.manuscript.ai_usage["request_count"], 1)
        self.assertEqual(self.manuscript.ai_consistency["status"], "clear")

    @patch("write.views.generate_draft")
    def test_failed_generate_does_not_deduct_tokens(self, generate_mock):
        generate_mock.side_effect = AIServiceError("AI request failed.")

        response = self.client.post(
            reverse("write:ai_generate", args=[self.manuscript.pk]),
            data=json.dumps({"action": "continue", "cursor_context": "river"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 502)
        wallet = AITokenWallet.objects.get(user=self.user)
        self.assertEqual(wallet.balance, 10000)
        self.assertEqual(wallet.total_used, 0)

    def test_zero_token_balance_returns_purchase_payload(self):
        settings = AITokenSettings.get_solo()
        settings.free_initial_tokens = 0
        settings.minimum_request_tokens = 50
        settings.save()

        response = self.client.post(
            reverse("write:ai_generate", args=[self.manuscript.pk]),
            data=json.dumps({"action": "continue", "cursor_context": "river"}),
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 402)
        self.assertEqual(response.json()["code"], "insufficient_tokens")
        self.assertIn("purchase_options", response.json())
