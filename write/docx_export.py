from io import BytesIO

from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_BREAK


def manuscript_docx_filename(manuscript):
    base = slugify(manuscript.title or "manuscript") or "manuscript"
    return f"{base}.docx"


def manuscript_to_docx_bytes(manuscript):
    document = Document()
    title = (manuscript.title or "").strip()
    if title:
        document.add_heading(title, level=0)

    _append_nodes(document, manuscript.content.get("content", []) if isinstance(manuscript.content, dict) else [])
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def submission_prefill_for(manuscript):
    text = _plain_text(manuscript.content)
    summary = " ".join(text.split())
    short = summary[:197].rstrip()
    if len(summary) > 197:
        short = f"{short}..."

    profile = manuscript.ai_profile if isinstance(manuscript.ai_profile, dict) else {}
    memory = manuscript.ai_memory if isinstance(manuscript.ai_memory, dict) else {}
    long_description = (
        memory.get("book_summary")
        or memory.get("summary")
        or summary[:1200]
        or f"Manuscript for {manuscript.title}."
    )

    return {
        "title": manuscript.title,
        "short_description": short or manuscript.title,
        "long_description": long_description,
        "language": _language_from_profile(profile, text),
        "category": _category_from_profile(profile),
    }


def _append_nodes(document, nodes, list_style=None):
    if not isinstance(nodes, list):
        return
    for node in nodes:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        content = node.get("content", [])
        if node_type == "heading":
            text = _inline_text(content).strip()
            if text:
                level = min(max(int((node.get("attrs") or {}).get("level") or 1), 1), 4)
                document.add_heading(text, level=level)
        elif node_type == "paragraph":
            paragraph = document.add_paragraph(style=list_style) if list_style else document.add_paragraph()
            _append_inline(paragraph, content)
        elif node_type == "blockquote":
            for child in content if isinstance(content, list) else []:
                paragraph = document.add_paragraph(style="Intense Quote")
                _append_inline(paragraph, child.get("content", []) if isinstance(child, dict) else [])
        elif node_type == "bulletList":
            _append_list(document, content, "List Bullet")
        elif node_type == "orderedList":
            _append_list(document, content, "List Number")
        elif node_type == "listItem":
            _append_nodes(document, content, list_style=list_style)
        elif content:
            _append_nodes(document, content, list_style=list_style)


def _append_list(document, items, style):
    for item in items if isinstance(items, list) else []:
        if not isinstance(item, dict):
            continue
        paragraphs = item.get("content", [])
        if not paragraphs:
            continue
        first = True
        for child in paragraphs:
            if not isinstance(child, dict):
                continue
            child_style = style if first and child.get("type") == "paragraph" else None
            _append_nodes(document, [child], list_style=child_style)
            first = False


def _append_inline(paragraph, nodes):
    for node in nodes if isinstance(nodes, list) else []:
        if not isinstance(node, dict):
            continue
        node_type = node.get("type")
        if node_type == "text":
            run = paragraph.add_run(node.get("text", ""))
            for mark in node.get("marks", []) or []:
                mark_type = mark.get("type") if isinstance(mark, dict) else ""
                if mark_type == "bold":
                    run.bold = True
                elif mark_type == "italic":
                    run.italic = True
                elif mark_type == "underline":
                    run.underline = True
                elif mark_type == "strike":
                    run.font.strike = True
        elif node_type == "hardBreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif node.get("content"):
            _append_inline(paragraph, node.get("content"))


def _inline_text(nodes):
    parts = []
    for node in nodes if isinstance(nodes, list) else []:
        if not isinstance(node, dict):
            continue
        if node.get("type") == "text":
            parts.append(node.get("text", ""))
        elif node.get("type") == "hardBreak":
            parts.append("\n")
        elif node.get("content"):
            parts.append(_inline_text(node.get("content")))
    return "".join(parts)


def _plain_text(content):
    if not isinstance(content, dict):
        return ""
    parts = []

    def walk(node):
        if not isinstance(node, dict):
            return
        node_type = node.get("type")
        if node_type == "text":
            parts.append(node.get("text", ""))
        elif node_type in {"paragraph", "heading", "listItem"}:
            for child in node.get("content", []) or []:
                walk(child)
            parts.append("\n")
        else:
            for child in node.get("content", []) or []:
                walk(child)

    walk(content)
    return "\n".join(line.strip() for line in "".join(parts).splitlines() if line.strip())


def _language_from_profile(profile, text):
    value = str(profile.get("language") or "").strip().lower()
    if value.startswith("fr") or "french" in value or "francais" in value or "français" in value:
        return "fr"
    if value.startswith("en") or "english" in value:
        return "en"
    french_markers = [" le ", " la ", " les ", " des ", " une ", " dans ", " pour ", " avec "]
    sample = f" {text[:2000].lower()} "
    return "fr" if sum(marker in sample for marker in french_markers) >= 3 else "en"


def _category_from_profile(profile):
    value = " ".join(str(profile.get(key) or "") for key in ("genre", "book_type")).lower()
    mapping = [
        ("romance", "romance"),
        ("thriller", "thriller_mystery"),
        ("mystery", "thriller_mystery"),
        ("drama", "drama"),
        ("science fiction", "scifi_fantasy"),
        ("fantasy", "scifi_fantasy"),
        ("horror", "horror"),
        ("poetry", "poetry"),
        ("biography", "biography_memoir"),
        ("memoir", "biography_memoir"),
        ("self", "self_help"),
        ("business", "business_money"),
        ("history", "history"),
        ("health", "health_wellness"),
        ("spiritual", "religion_spirituality"),
        ("religion", "religion_spirituality"),
        ("children", "children_ya"),
        ("academic", "academic"),
        ("course", "course"),
        ("politic", "politics"),
        ("african", "african_literature"),
        ("non", "non_fiction"),
    ]
    for needle, category in mapping:
        if needle in value:
            return category
    return "fiction"
