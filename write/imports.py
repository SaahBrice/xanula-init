from pathlib import Path
import re

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - exercised only when dependency is missing.
    PdfReader = None


MAX_IMPORT_BYTES = 10 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".docx", ".pdf"}


class ManuscriptImportError(ValueError):
    pass


def title_from_upload_name(name):
    stem = Path(name or "").stem.strip()
    stem = re.sub(r"[_-]+", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem or "Untitled"


def import_uploaded_manuscript(uploaded_file):
    extension = _validate_upload(uploaded_file)
    uploaded_file.seek(0)
    if extension == ".docx":
        content = _docx_to_tiptap(uploaded_file)
    elif extension == ".pdf":
        content = _pdf_to_tiptap(uploaded_file)
    else:
        raise ManuscriptImportError("Upload a DOCX or PDF file.")

    if not _content_has_text(content):
        raise ManuscriptImportError("No readable text was found in this file.")
    return content


def _validate_upload(uploaded_file):
    if not uploaded_file:
        raise ManuscriptImportError("Choose a DOCX or PDF file to upload.")
    size = int(getattr(uploaded_file, "size", 0) or 0)
    if size <= 0:
        raise ManuscriptImportError("The uploaded file is empty.")
    if size > MAX_IMPORT_BYTES:
        raise ManuscriptImportError("Upload a file smaller than 10 MB.")

    extension = Path(getattr(uploaded_file, "name", "") or "").suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ManuscriptImportError("Only DOCX and PDF files can be imported.")

    content_type = (getattr(uploaded_file, "content_type", "") or "").lower()
    allowed_content_types = {
        ".docx": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/octet-stream",
            "",
        },
        ".pdf": {
            "application/pdf",
            "application/octet-stream",
            "",
        },
    }
    if content_type not in allowed_content_types[extension]:
        raise ManuscriptImportError("The uploaded file type does not match its extension.")
    return extension


def _docx_to_tiptap(file_obj):
    try:
        from docx import Document
    except ImportError as exc:
        raise ManuscriptImportError("DOCX import is not configured yet.") from exc

    try:
        document = Document(file_obj)
    except Exception as exc:
        raise ManuscriptImportError("This DOCX file could not be read.") from exc

    nodes = []
    active_list = None
    active_list_type = None

    def flush_list():
        nonlocal active_list, active_list_type
        if active_list:
            nodes.append({"type": active_list_type, "content": active_list})
        active_list = None
        active_list_type = None

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            flush_list()
            continue

        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        list_type = _list_type_from_style(style_name)
        if list_type:
            if active_list_type != list_type:
                flush_list()
                active_list_type = list_type
                active_list = []
            active_list.append(_list_item_node(text))
            continue

        flush_list()
        heading_level = _heading_level_from_style(style_name)
        if heading_level:
            nodes.append(_heading_node(text, heading_level))
        else:
            nodes.append(_paragraph_node(text))

    flush_list()

    for table in document.tables:
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            text = " | ".join(cell for cell in cells if cell)
            if text:
                nodes.append(_paragraph_node(text))

    return {"type": "doc", "content": nodes}


def _pdf_to_tiptap(file_obj):
    if PdfReader is None:
        raise ManuscriptImportError("PDF import is not configured yet.")
    try:
        reader = PdfReader(file_obj)
    except Exception as exc:
        raise ManuscriptImportError("This PDF file could not be read.") from exc

    nodes = []
    for page in getattr(reader, "pages", []) or []:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        for paragraph in _split_pdf_paragraphs(page_text):
            nodes.append(_paragraph_node(paragraph))
    return {"type": "doc", "content": nodes}


def _split_pdf_paragraphs(text):
    text = _clean_text(text)
    if not text:
        return []
    parts = re.split(r"\n{2,}", text)
    if len(parts) <= 1:
        parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-Þ])", text)
    return [_clean_text(part) for part in parts if _clean_text(part)]


def _clean_text(text):
    text = re.sub(r"\r\n?", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]+", "\n", text)
    return text.strip()


def _heading_level_from_style(style_name):
    if style_name == "title":
        return 1
    match = re.match(r"heading\s+([1-6])", style_name or "")
    if match:
        return int(match.group(1))
    return 0


def _list_type_from_style(style_name):
    if "list" not in style_name:
        return ""
    if "number" in style_name:
        return "orderedList"
    return "bulletList"


def _text_content(text):
    return [{"type": "text", "text": text}]


def _paragraph_node(text):
    return {"type": "paragraph", "content": _text_content(text)}


def _heading_node(text, level):
    return {
        "type": "heading",
        "attrs": {"level": max(1, min(int(level or 1), 6))},
        "content": _text_content(text),
    }


def _list_item_node(text):
    return {"type": "listItem", "content": [_paragraph_node(text)]}


def _content_has_text(content):
    def walk(node):
        if isinstance(node, dict):
            if node.get("type") == "text" and str(node.get("text") or "").strip():
                return True
            return any(walk(child) for child in node.get("content", []) if isinstance(child, dict))
        if isinstance(node, list):
            return any(walk(child) for child in node)
        return False

    return walk(content)
